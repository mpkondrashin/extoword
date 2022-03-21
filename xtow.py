#!/usr/bin/env python
# -*- coding: utf-8 -*-
# cython: language_level=3
"""
    Convert RFP from Excel to Word.
"""

import sys
import os
from contextlib import contextmanager
from itertools import cycle
from time import time
from time import strftime

import xlrd
from docx import Document
from docx.shared import Mm
from docx.shared import Pt
from docx.shared import RGBColor

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell


def set_cell_border(cell: _Cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def single_border(cell):
    set_cell_border(
                cell,
                top={"sz": 4, "val": "single"},
                bottom={"sz": 4, "val": "single"},
                start={"sz": 4, "val": "single"},
                end={"sz": 4, "val": "single"},
            )


def head_cell_border(cell):
    set_cell_border(
                cell,
                top={"sz": 4, "val": "single"},
                bottom={"sz": 8, "val": "double"},
                start={"sz": 4, "val": "single"},
                end={"sz": 4, "val": "single"},
            )


def set_repeat_table_header(row):
    """ set repeat table row on every new page
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row


class Spin(object):

    minimal_iterval = 0.1
    start = ['(>      <)']*5
    transmission = ['(*      <)', '(>~     <)', '(> ~    <)', '(>  ~   <)', '(>   ~  <)', '(>    ~ <)', '(>     ~<)', '(>      *)']
    ph = []
    ph.extend(start)
    ph.extend(transmission)
    ph.extend(start)
    ph.extend(reversed(transmission))

    phases = [
        ['bq', 'dp', 'qb', 'pd'],
        'TYVY',
        ',.!.',
        ph,
        ['/ ', '| ', '\\ ', '--', ' /', ' |', ' \\', ' |', ' /', '--', '\\ ', '| '],
        ['/ ', '| ', '\\ ', '--', ' /', ' |', ' \\', '--'],
        '.oOo.',
        ['>====', '=>===', '==>==', '===>=', '====>', '====v', '====<', '===<=',
         '==<==', '=<===', '<====', '^===='],
         '>v<^',
         '+x',
         '.-+*+-',
         '-+-x',
         "_-`'-",
         '_\\|/_',
         ",-'-",
         '/-\\|',
    ]
    phases_count = len(phases)
    def __init__(self, style: int=0, stream=sys.stdout):
        self.stream = stream
        self.style = style
        self.phase_iter = cycle(self.phases[style])
        self.last_show = None
        #self.write()

    def __call__(self):
        self.advance()

    def advance(self):
        if self.last_show is None:
            self.cr()
            self.write()
            self.last_show = time()
            return
        t = time()
        if t - self.last_show < self.minimal_iterval:
            return
        self.last_show = t
        self.cr()
        self.write()

    def write(self):
        text = next(self.phase_iter)
        self.stream.write(text)
        self.stream.flush()

    def size(self):
        return len(self.phases[self.style][0])

    def placeholder(self):
        return ' ' * self.size()

    def cr(self):
        self.stream.write('\b' * self.size())

    def cleanup(self):
        self.cr()
        self.stream.write(' ' * self.size())
        self.cr()


@contextmanager
def spinner(style=0):
    s = Spin(style)
    yield s
    s.cleanup()


class CellStyle():
    """
        ...
        'Heading 1': (0, 49), 'Heading 2': (0, 50), 'Heading 3': (0, 51),
        'Heading 4': (0, 52), 'Input': (0, 53), 'Linked Cell': (0, 54),
        ...

    """
    def __init__(self, wb):
        self.wb = wb
        self.mapping = {index: style_name for style_name, (_, index) in wb.style_name_map.items()}

    def __call__(self, cell):
        key = self.wb.xf_list[cell.xf_index].parent_style_index
        return self.mapping[key]


def map_style_to_heading_size(style: str) -> int:
    style_to_heading_size_mapping = {
        'Heading 1': 0,
        'Heading 2': 1,
        'Heading 3': 2,
        'Heading 4': 3,
        'Заголовок 1': 0,
        'Заголовок 2': 1,
        'Заголовок 3': 2,
        'Заголовок 4': 3,
    }
    try:
        return style_to_heading_size_mapping[style]
    except KeyError:
        return None


def text_and_list_style(text: str):
    if text[0] in ['-', '—']:
        text = text[1:].strip()
        if text[0] in ['-', '—']:
            text = text[1:].strip()
            return text, 'List Bullet 3'
        else:
            return text, 'List Bullet 2'
    else:
        return text, 'List Number'


def item_depth(text: str):
    if text[0] in ['-', '—']:
        text = text[1:].strip()
        if text[0] in ['-', '—']:
            text = text[1:].strip()
            return text, 2
        else:
            return text, 1
    else:
        return text, 0


def default_template():
    file_name = 'default.docx'
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, 'media', file_name)
    else:
        return None


class ExcelToWord(object):

    def __init__(self,
                 input_files: list,
                 file_name_docx: str,
                 criteria: list=list(),
                 sheets: list=list()):
        self.input_files = input_files
        self.file_name_docx = file_name_docx
        #self.wb = xlrd.open_workbook(file_name_xls, formatting_info=True)
        self.wb = [xlrd.open_workbook(file_name_xls, formatting_info=True)
                   for file_name_xls in input_files]
        self.style = None

        self.criteria = criteria
        self.sheets = sheets
        # Hint for pyinstaller to include default.docx into package
        # './venv/lib/python3.7/site-packages/docx/templates/default.docx'
        self.doc = Document(default_template())
        self.count_requirements = 0

    def criteria_dict(self, sheet):
        result = dict()
        for n in range(2, sheet.ncols):
            cell = sheet.cell(0, n)
            value = cell.value.strip()
            if value == '':
                break
            if value in result:
                raise RuntimeError('Criteria names should be unique')
            result[value] = n
        return result

    """
    def list_packages(self):
        result = []
        for w in self.wb:
            book_result = list()
            for sheet in w.sheets()[1:]:
                for key in self.criteria_dict(sheet).keys():
                    if key in book_result:
                        continue
                    book_result.append(key)
            result.append(book_result)
        return result
    """
    def list_packages(self):
        result = []
        for w in self.wb:
            for sheet in w.sheets()[1:]:
                for key in self.criteria_dict(sheet).keys():
                    if key in result:
                        continue
                    result.append(key)
            #result.append(book_result)
        return result

    def list_sheets(self):
        result = []
        for w in self.wb:
            for sheet in w.sheets()[1:]:
                result.append(sheet.name)
            #result.append(book_result)
        return result

    def prefix(self):
        # self.add_toc()
        pass

    def add_toc(self):
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')  # creates a new element
        fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
        instrText.text = 'TOC \o "1-3" \h \z \\u'  # change 1-3 depending on heading levels you need

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:t')
        fldChar3.text = "Right-click to update field."
        fldChar2.append(fldChar3)

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        r_element = run._r
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar4)
        p_element = paragraph._p

    def heading(self, text: str, level: int):
        pass

    def item(self, text: str, depth: int):
        pass

    def count_scope_lines(self):
        result = 0
        for w in self.wb:
            for sheet in w.sheets()[1:]:
                result += (sheet.nrows-1)
        return result

    def run(self):
        self.prefix()
        yield from self.process()
        self.footer()

    def checked(self, value):
        if isinstance(value, str):
            return value.strip() != ''
        elif isinstance(value, int):
            return value != 0
        elif isinstance(value, float):
            return value != 0.0
        else:
            raise RuntimeError('{!r}: Wrond option cell type'.format(value))

    def should_skip(self, row, criteria_heading):
#        print('should_skip({}, {})'.format('|'.join([str(r.value) for r in row]), criteria_heading))
        if not self.criteria:
#            print('Skip')
            return False
        # criteria = set(SPE, TMDS-AM, DDAN)
        # criteria_heading = {SPE: 3, SPC: 4, ESEL: 5}
        for crit in self.criteria:
#            print('check crit = |{}|'.format(crit))
            if crit not in criteria_heading:
#                print('not in criteria_heading')
                continue
            col = criteria_heading[crit]
#            print('col = {}'.format(col))
            if self.checked(row[col].value):
#                print('is checked')
                return False
        return True

    def process(self):
        count = 0
        for w, file_name in zip(self.wb, self.input_files):
            self.style = CellStyle(w)
            for n, sheet in enumerate(w.sheets()[1:]):
                if self.sheets and sheet.name not in self.sheets:
                    print('{} not in {}'.format(sheet.name, self.sheets))
                    continue
                criteria_heading = self.criteria_dict(sheet)
    #            print('%s: ' % sheet.name, end='')
                #from random import randint
    #            with spinner(n) as spin:
                for line in range(1, sheet.nrows):
                    #spin()
                    if self.should_skip(sheet.row(line), criteria_heading):
#                        print('Skip!')
                        continue
                    count += 1
                    yield os.path.basename(file_name), sheet.name, count
#                    print('Do not skip')
                    cell = sheet.cell(line, 1)
                    if cell.value is None:
                        break
                    try:
                        text = cell.value.strip()
                    except AttributeError:
                        raise RuntimeError(f'Error on line {line+1} on sheet "{sheet.name}"')
                    if text == '':
                        continue
                    heading = map_style_to_heading_size(self.style(cell))
                    if heading is None:
                        self.count_requirements += 1
                        text, depth = item_depth(text)
                        self.item(text, depth)
                    else:
                        self.heading(text, level=heading)
    #            print('done')

    def footer(self):
        today = strftime('%d%m%Y')
        date_marker = 'GD{}DG'.format(today)
        p = self.doc.add_paragraph(date_marker)
        p.style = self.doc.styles['Body Text']
        p.style.font.size = Pt(2)
        p.style.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        self.doc.save(self.file_name_docx)


class Counter(ExcelToWord):
    name = 'counter'

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.counter = 0

    def heading(self, text: str, level: int):
        self.counter += 1

    def item(self, text: str, depth: int):
        self.counter += 1

    def count_scope_lines(self):
        return sum((1 for _ in self.process()))

class ExcelToWordList(ExcelToWord):

    name = 'list'

    def heading(self, text: str, level: int):
        self.doc.add_heading(text, level=level)

    def item(self, text: str, depth: int):
        style = ['List Number', 'List Bullet 2', 'List Bullet 3']
        self.doc.add_paragraph(text, style=style[depth])


MAX_DEPTH = 6


class ExcelToWordTable(ExcelToWord):

    name = 'fancy'

    widths = (Mm(16), Mm(120), Mm(12))

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.table = None
        self.counters = [0]*MAX_DEPTH
        self.last_heading_depth = 1

    def setup_widths(self, row):
        for cell, width in zip(row.cells, self.widths):
            cell.width = width

    def prefix(self):
        #self.table = self.doc.add_table(rows=1, cols=3, style='Medium Grid 1 Accent 3')
        self.table = self.doc.add_table(rows=1, cols=3)
#        self.table.autofit = True
        row = self.table.rows[0]
        self.setup_widths(row)
        set_repeat_table_header(row)
        row.cells[0].text = '№ пп'
        row.cells[1].text = 'Требование'
        row.cells[2].text = 'Да/Нет'
        for c in row.cells:
            single_border(c)

    def index(self, level: int) -> str:
        """ return str index '1.3.2' """
        if level == 0:
            return ''
        level -= 1
        self.counters[level] += 1
        for i in range(level + 1, MAX_DEPTH):
            self.counters[i] = 0
        return '.'.join([str(c) for c in self.counters[0:level+1]])

    def heading(self, text: str, level: int):
        self.last_heading_depth = level+1
        row_cells = self.table.add_row().cells
        row_cells[0].merge(row_cells[2])
        p = row_cells[0].paragraphs[0]
        p.style = 'Heading %s' % (level + 1)
        p.text = '{}  {}'.format(self.index(level), text)

    def item(self, text: str, depth: int):
        row = self.table.add_row()
        self.setup_widths(row)
        row.cells[0].text = self.index(self.last_heading_depth + depth)
        style = [None, 'List Bullet', 'List Bullet 2']
        #style = [None, 'List Bullet 2', 'List Bullet 3']
        p = row.cells[1].paragraphs[0]
        p.style = style[depth]
        p.text = text
        single_border(row.cells[2])
        row.cells[2].text = ' '


class ExcelToWordPlainTable(ExcelToWordTable):

    name = 'table'

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.table = None
        self.counters = [0]*MAX_DEPTH
        self.last_heading_depth = 1

    def prefix(self):
        self.table = self.doc.add_table(rows=1, cols=3,
                                        style='Table Grid')

        self.setup_widths(self.table.rows[0])
        set_repeat_table_header(self.table.rows[0])
        hdr_cells = self.table.rows[0].cells
        hdr_cells[0].text = '№ пп'
        hdr_cells[1].text = 'Требование'
        hdr_cells[2].text = 'Да/Нет'
        for cell in hdr_cells:
            head_cell_border(cell)

    def heading(self, text: str, level: int):
        self.last_heading_depth = level+1
        row = self.table.add_row()
        self.setup_widths(row)
        row_cells = row.cells
        row_cells[0].add_paragraph(self.index(level))
        row_cells[1].add_paragraph(text)
        row_cells[2].text = ''

requirements_ru = """

Этот текст уже не актуален! См. английскую версию

Требования к формату файла MS Excel:

- Формат файла MS Excel xls. (xlsx не поддерживается!)
- Предметная информация размещается в листах, начиная со второго (первый
  игнорируется конвертором)
- Разбиение на листы — только для удобства. Конвертер проходит по ним
  последовательно слева направо
- На каждом листе первая колонка зарезервирована для будущих параметров
  конвертора
- Все требования размещаются во второй колонке (прочие колонки игнорируются
  конвертором)
- Третья и последующие колонки могут быть использованы для формирования
  выборок
- Пункты списка отмечаются символом '-' (дефис) или '—' (тире)
- Для пунктов второго уровня вложенности, символы '-' или '—' дублируются
- Для разбиения документа на разделы могут использоваться заголовки стиля
  Heading с первого по четвертый.
- Заголовки прочих размеров и любое другое форматирование — игнорируются
"""

description='Convert RFP MS Excel file to MS Word'

requirements_en = """MS Excel file format requirements:

— MS Excel file should be xls and not xlsx
— Excel book sheets are processed starting from second one (first sheet is ignored by converter)
— Splitting list of requirements to separate sheets is only for convenience. Converter processes them from left to right consecutively
— First column on each sheet is reserved for converter parameters
— All requirements are placed into second column
— Third and other columns can be used for checkmarks for selection of subsets of requirements
— First line on each sheet is for subset names (i.e. software packages)
— List of items are marked by '-' (hyphen) or '—' (dash) in first position
— Double symbols '-' or '—' mark second level list items
— To separate list of requirements styles Heading one to four can be used
— All other style headings and any other formatting are ignored
"""


def parse_args():
    import argparse
    parser = argparse.ArgumentParser(description=description,
                                     epilog=requirements_en,
                                     formatter_class = argparse.RawDescriptionHelpFormatter
    )
    parser.add_argument('-f', '--format',
                        choices=[ExcelToWordList.name,
                                 ExcelToWordPlainTable.name,
                                 ExcelToWordTable.name],
                        help='Output file format')
    parser.add_argument("files", nargs='+', help="xls filename to convert")
    return parser.parse_args()


def output_file_name(input_file_name, folder=''):
    name, ext = os.path.splitext(input_file_name)
    today = strftime('%d%m%Y')
    return os.path.join(folder, '{}_{}.docx'.format(name, today))


@contextmanager
def timer(fmt='Process time: {} sec'):
    time_start = time()
    yield
    time_end = time()
    print(fmt.format(int(time_end - time_start)))


def convert(converter):
    current_phase = None
    with timer():
        with spinner(5) as spin:
            for file_name, sheet, line in converter.run():
                phase = '%s: %s' % (file_name, sheet)
                if current_phase is None:
                    current_phase = phase
                    print('%s: %s' % (current_phase, spin.placeholder()), end='')
                elif phase != current_phase:
                    current_phase = phase
                    spin.cleanup()
                    print('done\n%s: %s' % (current_phase, spin.placeholder()), end='')
                spin()
    print('Total number of requirements: {}'.format(converter.count_requirements))
    print('Output written to: {}'.format(output_file_name_docx))


available_converters = {
    ExcelToWordTable.name: ExcelToWordTable,
    ExcelToWordPlainTable.name: ExcelToWordPlainTable,
    ExcelToWordList.name: ExcelToWordList
}

default_converter = ExcelToWordList

if __name__ == '__main__':
    args = parse_args()
    output_file_name_docx = output_file_name(args.files[0])
    try:
        converter_class = available_converters[args.format]
    except KeyError:
        converter_class = ExcelToWordList

    converter = converter_class(args.files, output_file_name_docx)
    print(converter.list_packages())
    convert(converter)
